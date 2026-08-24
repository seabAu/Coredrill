"""Generate deterministic, synthetic DOCX/PDF fixtures for local import tests."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from datetime import datetime, timezone
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas


REPOSITORY_ROOT = Path(__file__).resolve().parents[2]
FIXTURE_ROOT = REPOSITORY_ROOT / "fixtures" / "imports"
DOCX_PATH = FIXTURE_ROOT / "synthetic-resume.docx"
PDF_PATH = FIXTURE_ROOT / "synthetic-two-page.pdf"
SCANNED_PDF_PATH = FIXTURE_ROOT / "synthetic-scanned.pdf"
FIXED_TIMESTAMP = datetime(2026, 8, 24, 12, 0, 0, tzinfo=timezone.utc)


def normalize_docx_archive(path: Path) -> None:
    with ZipFile(path, "r") as source:
        entries = [(info.filename, source.read(info.filename)) for info in source.infolist()]
    with ZipFile(path, "w", compression=ZIP_DEFLATED, compresslevel=9) as target:
        for filename, data in sorted(entries):
            info = ZipInfo(filename, date_time=(2026, 8, 24, 12, 0, 0))
            info.compress_type = ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            target.writestr(info, data)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "145DA0")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend([properties, text_element])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 22, "12344D"),
        ("Heading 1", 13, "12344D"),
        ("Heading 2", 11, "145DA0"),
        ("Heading 3", 10.5, "145DA0"),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(3)

    title = document.add_paragraph(style="Title")
    title.add_run("Jordan Rivera")
    contact = document.add_paragraph()
    contact.paragraph_format.space_after = Pt(8)
    contact.add_run("Product Engineer · jordan@example.test · ")
    add_hyperlink(contact, "Portfolio", "https://example.test/portfolio")

    document.add_paragraph("Summary", style="Heading 1")
    summary = document.add_paragraph()
    summary.add_run("Built an ")
    accessible = summary.add_run("accessible")
    accessible.bold = True
    summary.add_run(" local-first workflow for ")
    evidence = summary.add_run("evidence-backed")
    evidence.italic = True
    summary.add_run(" job decisions.")

    document.add_paragraph("Highlights", style="Heading 2")
    document.add_paragraph(
        "Reduced research turnaround by 42% with explicit source trails.",
        style="List Bullet",
    )
    document.add_paragraph(
        "Kept user-confirmed facts separate from imported proposals.",
        style="List Bullet",
    )
    document.add_paragraph("Experience", style="Heading 1")
    document.add_paragraph("Coredrill Labs — Product Engineer — 2024–2026")

    document.core_properties.title = "Synthetic resume import fixture"
    document.core_properties.author = "Coredrill test suite"
    document.core_properties.created = FIXED_TIMESTAMP
    document.core_properties.modified = FIXED_TIMESTAMP
    document.save(DOCX_PATH)
    normalize_docx_archive(DOCX_PATH)


def create_text_pdf() -> None:
    pdf = canvas.Canvas(str(PDF_PATH), pagesize=LETTER, invariant=1, pageCompression=1)
    pdf.setTitle("Synthetic two-page import fixture")
    pdf.setAuthor("Coredrill test suite")
    pdf.setFont("Helvetica-Bold", 18)
    pdf.drawString(54, 738, "Coredrill PDF Fixture")
    pdf.setFont("Helvetica", 11)
    pdf.drawString(54, 708, "Page one evidence line.")
    pdf.drawString(54, 690, "Source-backed decisions remain proposals until confirmed.")
    pdf.showPage()
    pdf.setFont("Helvetica-Bold", 15)
    pdf.drawString(54, 738, "Salary Notes")
    pdf.setFont("Helvetica", 11)
    pdf.drawString(54, 708, "Page two salary source line.")
    pdf.save()


def create_scanned_pdf() -> None:
    image = Image.new("RGB", (1275, 1650), "white")
    drawing = ImageDraw.Draw(image)
    drawing.rectangle((70, 70, 1205, 1580), outline="#12344d", width=5)
    drawing.text((120, 150), "SCANNED RESUME IMAGE", fill="#12344d")
    drawing.text((120, 220), "This visible text is pixels, not extractable PDF text.", fill="black")
    image_bytes = BytesIO()
    image.save(image_bytes, format="PNG", optimize=False)
    image_bytes.seek(0)

    pdf = canvas.Canvas(str(SCANNED_PDF_PATH), pagesize=LETTER, invariant=1, pageCompression=1)
    pdf.setTitle("Synthetic scanned import fixture")
    pdf.drawImage(ImageReader(image_bytes), 0, 0, width=612, height=792, mask="auto")
    pdf.save()


def main() -> None:
    FIXTURE_ROOT.mkdir(parents=True, exist_ok=True)
    create_docx()
    create_text_pdf()
    create_scanned_pdf()
    for path in (DOCX_PATH, PDF_PATH, SCANNED_PDF_PATH):
        print(f"generated {path.relative_to(REPOSITORY_ROOT)} ({path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
